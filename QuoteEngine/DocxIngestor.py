import docx 
from QuoteEngine.IngestorInterface import IngestorInterface
    

class DocxIngestor(IngestorInterface):
    def __init__(self, path):
        '''Takes a docx file and processes the contents into a list.'''
        self.path = path
        
    def get_quotes(path, file_type):
        '''Returns a list with the quote str and author str.'''
        if IngestorInterface.parse(path):
            text_list = []
            document = docx.Document(path)
            for x in range(len(document.paragraphs)):
                text = document.paragraphs[x].text
                text_list.append(text)
            return text_list
          
    
